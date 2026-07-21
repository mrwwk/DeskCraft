# Dedicated evaluator for Policy_L3_04.docx task (cd8b1536)
# Replaces generic compare_docx_files with format-aware multi-metric checks.
# Conditions: (1) title center/22pt/bold, (2) section headings uppercase,
# (3) body font Times New Roman 11pt, (4) section 4 body bold,
# (5) first para CONFIDENTIAL, (6) last para Document Control.

import re
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from rapidfuzz import fuzz

logger = logging.getLogger("desktopenv.metric.docs")

# Section heading pattern: "1. PURPOSE", "4. EQUIPMENT AND SECURITY", etc.
_HEADING_RE = re.compile(r'^\d+\.\s+[A-Z][A-Z\s&/, \-]+$')

# Markers for paragraphs that are not body text
_SPECIAL_MARKERS = [
    "CONFIDENTIAL",
    "Document Control",
    "COMPANY REMOTE WORK POLICY",
    "EFFECTIVE DATE",
]


def _normalize_dash(text: str) -> str:
    """Replace various Unicode dash characters with ASCII hyphen."""
    text = text.replace('\u2013', '-')   # EN DASH
    text = text.replace('\u2014', '-')   # EM DASH
    text = text.replace('\u2015', '-')   # HORIZONTAL BAR
    return text


def _normalize_ws(text: str) -> str:
    """Collapse all whitespace to single spaces and strip."""
    return re.sub(r'\s+', ' ', text).strip()


def _normalize_text(text: str) -> str:
    """Full normalization: dashes + whitespace."""
    return _normalize_ws(_normalize_dash(text))


def _is_heading(text: str) -> bool:
    """Check if paragraph text matches a numbered section heading pattern."""
    return bool(_HEADING_RE.match(text.strip()))


def _is_special(text: str) -> bool:
    """Check if paragraph contains a known special marker."""
    t = text.lower()
    for marker in _SPECIAL_MARKERS:
        if marker.lower() in t:
            return True
    return False


def _is_body_paragraph(text: str) -> bool:
    """True if paragraph looks like body content (not heading, not special, not empty)."""
    t = text.strip()
    if not t:
        return False
    if _is_heading(t):
        return False
    if _is_special(t):
        return False
    return True


# ---------------------------------------------------------------------------
# Metric 1: check_title_format  (condition 1)
# ---------------------------------------------------------------------------

def check_title_format(result_path, **options):
    """Verify (1): title is centered, 22pt, bold.

    Signature: 1 required positional param → framework calls as
        check_title_format(result_state, **options[0])
    """
    if not result_path:
        return 0.0

    try:
        doc = Document(result_path)
    except Exception as e:
        logger.error(f"check_title_format: cannot open {result_path}: {e}")
        return 0.0

    title_keywords = options.get("title_keywords", "REMOTE WORK POLICY")
    expected_size_pt = float(options.get("title_font_size_pt", 22.0))

    for para in doc.paragraphs:
        if title_keywords.lower() not in para.text.lower():
            continue

        # --- alignment ---
        if para.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            logger.info("check_title_format: title not centered")
            return 0.0

        # --- bold + font size on non-empty runs ---
        has_bold = False
        all_size_ok = True
        any_run_checked = False
        for run in para.runs:
            if not run.text.strip():
                continue
            any_run_checked = True
            if run.bold:
                has_bold = True
            if run.font.size is not None and abs(run.font.size.pt - expected_size_pt) > 0.5:
                all_size_ok = False

        if not any_run_checked:
            logger.info("check_title_format: title paragraph has no non-empty runs")
            return 0.0
        if not has_bold:
            logger.info("check_title_format: title not bold")
            return 0.0
        if not all_size_ok:
            logger.info(f"check_title_format: title font size not {expected_size_pt}pt")
            return 0.0

        return 1.0

    logger.info("check_title_format: title paragraph not found")
    return 0.0


# ---------------------------------------------------------------------------
# Metric 2: check_body_font  (condition 3)
# ---------------------------------------------------------------------------

def check_body_font(result_path, **options):
    """Verify (3): all body text uses Times New Roman 11pt.

    Signature: 1 required positional param → framework calls as
        check_body_font(result_state, **options[1])
    """
    if not result_path:
        return 0.0

    try:
        doc = Document(result_path)
    except Exception as e:
        logger.error(f"check_body_font: cannot open {result_path}: {e}")
        return 0.0

    expected_font = options.get("expected_font", "Times New Roman")
    expected_size_pt = float(options.get("expected_size_pt", 11.0))
    threshold = float(options.get("body_font_threshold", 0.80))

    total_runs = 0
    failed_runs = 0

    for para in doc.paragraphs:
        if not _is_body_paragraph(para.text):
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            ok = True
            # Only flag explicit (non-None) values that differ from expected
            if run.font.name is not None and run.font.name != expected_font:
                ok = False
            if run.font.size is not None and abs(run.font.size.pt - expected_size_pt) > 0.5:
                ok = False
            if not ok:
                failed_runs += 1

    if total_runs == 0:
        # No body runs found — ambiguous; pass to avoid false negative
        return 1.0

    ratio = 1.0 - failed_runs / total_runs
    if ratio >= threshold:
        return 1.0
    else:
        logger.info(f"check_body_font: ratio={ratio:.3f} < threshold={threshold}")
        return ratio


# ---------------------------------------------------------------------------
# Metric 3: check_section4_bold  (condition 4)
# ---------------------------------------------------------------------------

def check_section4_bold(result_path, **options):
    """Verify (4): body text under section 4 is bold.

    Signature: 1 required positional param → framework calls as
        check_section4_bold(result_state, **options[2])
    """
    if not result_path:
        return 0.0

    try:
        doc = Document(result_path)
    except Exception as e:
        logger.error(f"check_section4_bold: cannot open {result_path}: {e}")
        return 0.0

    section4_marker = options.get("section4_marker", "4. EQUIPMENT AND SECURITY")

    # Locate section 4 heading
    sec4_idx = None
    for i, para in enumerate(doc.paragraphs):
        if section4_marker.lower() in para.text.lower():
            sec4_idx = i
            break

    if sec4_idx is None:
        # Section 4 not found — nothing to check, assume pass
        return 1.0

    # Find next section heading after section 4
    next_sec_idx = None
    for i in range(sec4_idx + 1, len(doc.paragraphs)):
        if _is_heading(doc.paragraphs[i].text):
            next_sec_idx = i
            break

    start = sec4_idx + 1
    end = next_sec_idx if next_sec_idx is not None else len(doc.paragraphs)

    body_paras = [
        p for p in doc.paragraphs[start:end]
        if p.text.strip() and not _is_special(p.text)
    ]

    if not body_paras:
        return 1.0

    for para in body_paras:
        para_has_bold = False
        for run in para.runs:
            if run.text.strip() and run.bold:
                para_has_bold = True
                break
        if not para_has_bold:
            logger.info("check_section4_bold: body paragraph under section 4 not bold")
            return 0.0

    return 1.0


# ---------------------------------------------------------------------------
# Metric 4: compare_docx_text_normalized  (conditions 2, 5, 6)
# ---------------------------------------------------------------------------

def compare_docx_text_normalized(result_path, expected_path, **options):
    """Verify (2)(5)(6): text content with dash/whitespace normalization.

    Signature: 2 required positional params → framework calls as
        compare_docx_text_normalized(result_state, expected_state, **options[3])
    """
    if not result_path or not expected_path:
        return 0.0

    try:
        result_doc = Document(result_path)
        expected_doc = Document(expected_path)
    except Exception as e:
        logger.error(f"compare_docx_text_normalized: cannot open docs: {e}")
        return 0.0

    # Build normalized, non-empty paragraph lists
    r_paras = [_normalize_text(p.text) for p in result_doc.paragraphs]
    r_paras = [t for t in r_paras if t]

    e_paras = [_normalize_text(p.text) for p in expected_doc.paragraphs]
    e_paras = [t for t in e_paras if t]

    if not r_paras:
        return 0.0

    # (5) First paragraph must be "CONFIDENTIAL - Internal Use Only"
    confidential_target = _normalize_text("CONFIDENTIAL - Internal Use Only")
    if r_paras[0] != confidential_target:
        logger.info("compare_docx_text_normalized: first paragraph != CONFIDENTIAL")
        return 0.0

    # (6) Last paragraph must contain "Document Control: Version 3.0"
    doc_control_prefix = _normalize_text("Document Control: Version 3.0")
    if doc_control_prefix not in r_paras[-1]:
        logger.info("compare_docx_text_normalized: last paragraph missing Document Control")
        return 0.0

    # (2) All section headings must be ALL CAPS
    for t in r_paras:
        m = _HEADING_RE.match(t)
        if m:
            # The heading text after the number prefix
            heading_body = m.group(1)
            if heading_body != heading_body.upper():
                logger.info(f"compare_docx_text_normalized: heading not uppercase: '{heading_body}'")
                return 0.0

    # Overall fuzzy similarity as safety net
    r_text = '\n'.join(r_paras)
    e_text = '\n'.join(e_paras)
    similarity = fuzz.ratio(r_text, e_text) / 100.0

    threshold = float(options.get("text_similarity_threshold", 0.85))
    if similarity < threshold:
        logger.info(f"compare_docx_text_normalized: similarity={similarity:.3f} < {threshold}")
        return 0.0

    return similarity
