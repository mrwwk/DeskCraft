"""
Evaluator for L3_writ task: format and edit Guide_L3_05.docx in LibreOffice Writer.

Metrics:
  0. compare_docx_files       – text content comparison (requirements 1, 6)
  1. check_title_formatting   – title centered, 24pt bold (requirement 2)
  2. check_section_heading_format – section headings bold + dark green #006400 (requirement 3)
  3. check_body_text_font_size    – body text 11pt (requirement 4)
  4. check_week4_body_italic      – Week 4 body italic (requirement 5)
"""
import re
import logging

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

logger = logging.getLogger("desktopenv.metric.docs")


# ── Metric 0: text content comparison ──────────────────────────────────────

def compare_docx_files(file1, file2, **options):
    """Compare text content of two .docx files, normalising whitespace.

    This covers task requirements (1) replace text and (6) add Questions line,
    without checking formatting attributes.
    """
    ignore_blanks = options.get('ignore_blanks', True)
    ignore_case = options.get('ignore_case', False)

    if not file1 or not file2:
        return 0

    try:
        doc1 = Document(file1)
        doc2 = Document(file2)
    except Exception:
        logger.error("Failed to open one or both .docx files", exc_info=True)
        return 0

    paragraphs1 = [p.text for p in doc1.paragraphs]
    paragraphs2 = [p.text for p in doc2.paragraphs]

    if ignore_blanks:
        text1 = re.sub(r'\s+', ' ', '\n'.join(paragraphs1)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(paragraphs2)).strip()
        if ignore_case:
            text1 = text1.lower()
            text2 = text2.lower()
        if text1 != text2:
            return 0
    else:
        if len(paragraphs1) != len(paragraphs2):
            return 0
        for p1, p2 in zip(paragraphs1, paragraphs2):
            if ignore_case:
                p1, p2 = p1.lower(), p2.lower()
            if p1 != p2:
                return 0

    return 1


# ── Metric 1: title formatting (centered, 24pt, bold) ──────────────────────

def check_title_formatting(result_file, expected=None, **options):
    """Verify the first paragraph (title) is CENTER aligned, 24pt, and bold."""
    if not result_file:
        return 0

    try:
        doc = Document(result_file)
    except Exception:
        logger.error("Failed to open result file", exc_info=True)
        return 0

    if not doc.paragraphs:
        return 0

    first_para = doc.paragraphs[0]

    # Check centered alignment
    if first_para.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
        return 0

    # Check at least one non-empty run is bold AND 24pt
    for run in first_para.runs:
        if run.text.strip() and run.bold:
            if run.font.size and run.font.size.pt == 24:
                return 1

    return 0


# ── Metric 2: section heading formatting (bold + dark green) ───────────────

def check_section_heading_format(result_file, expected=None, **options):
    """All paragraphs starting with 'Week <N>' must be bold and dark green (#006400)."""
    if not result_file:
        return 0

    try:
        doc = Document(result_file)
    except Exception:
        logger.error("Failed to open result file", exc_info=True)
        return 0

    target_color = RGBColor(0x00, 0x64, 0x00)
    heading_pattern = re.compile(r'^Week\s+\d')

    found_any_heading = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not heading_pattern.match(text):
            continue
        found_any_heading = True

        # Check that at least one run in this heading paragraph is bold + dark green
        found_bold_green = False
        for run in para.runs:
            if not run.text.strip():
                continue
            if run.bold and run.font.color and run.font.color.rgb == target_color:
                found_bold_green = True
                break
        if not found_bold_green:
            return 0

    # If no headings found at all, something is wrong → fail
    if not found_any_heading:
        return 0

    return 1


# ── Metric 3: body text font size 11pt ─────────────────────────────────────

def check_body_text_font_size(result_file, expected=None, **options):
    """All non-title, non-heading paragraphs must have 11pt font size."""
    if not result_file:
        return 0

    try:
        doc = Document(result_file)
    except Exception:
        logger.error("Failed to open result file", exc_info=True)
        return 0

    heading_pattern = re.compile(r'^Week\s+\d')

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Skip title (paragraph 0)
        if i == 0:
            continue
        # Skip section headings
        if heading_pattern.match(text):
            continue
        # Skip truly empty paragraphs (no text in any run)
        if not text:
            continue

        # Check every run with visible text has font size 11pt
        for run in para.runs:
            if not run.text.strip():
                continue
            # If font.size is None, the run may inherit from paragraph style;
            # be lenient here and only fail when an explicit wrong size is set.
            if run.font.size is not None and run.font.size.pt != 11:
                return 0

    return 1


# ── Metric 4: Week 4 body italic ───────────────────────────────────────────

def check_week4_body_italic(result_file, expected=None, **options):
    """Body paragraphs immediately after 'Week 4: Integration' must be italic."""
    if not result_file:
        return 0

    try:
        doc = Document(result_file)
    except Exception:
        logger.error("Failed to open result file", exc_info=True)
        return 0

    # Locate "Week 4" heading
    week4_idx = None
    for i, para in enumerate(doc.paragraphs):
        if 'Week 4' in para.text and 'Integration' in para.text:
            week4_idx = i
            break

    if week4_idx is None:
        # Week 4 heading not found — this metric passes vacuously
        # (text content metric will catch missing sections)
        return 1

    heading_pattern = re.compile(r'^Week\s+\d')
    has_body = False

    # Check every body paragraph between Week 4 heading and next heading (or EOF)
    for i in range(week4_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if heading_pattern.match(text):
            break  # next section heading reached
        if not text:
            continue
        has_body = True

        # At least one non-empty run must be italic
        any_italic = False
        for run in doc.paragraphs[i].runs:
            if run.text.strip() and run.italic:
                any_italic = True
                break
        if not any_italic:
            return 0

    # If there is no body text after Week 4 heading, this is suspicious
    # but let text comparison catch it; pass vacuously here.
    if not has_body:
        return 1

    return 1
