import logging
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger("desktopenv.metric.docs")


# =============================================================================
# Metric 0: compare_docx_files — 覆盖条件(2)Q1/Q2替换 和 条件(6)结束标记
# 通过段落文本比对间接验证文本替换和结束标记
# =============================================================================

def compare_docx_files(file1, file2, **options):
    ignore_blanks = options.get('ignore_blanks', True)
    ignore_case = options.get('ignore_case', False)
    ignore_order = options.get('ignore_order', False)
    content_only = options.get('content_only', False)
    fuzzy_match = options.get('fuzzy_match', False)
    delete_empty_lines = options.get('delete_empty_lines', False)

    if not file1 or not file2:
        return 0

    # Determine file types and load documents
    if file1.endswith('.docx') and file2.endswith('.docx'):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0
        doc1_paragraphs = [p.text for p in doc1.paragraphs]
        doc2_paragraphs = [p.text for p in doc2.paragraphs]
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
        if delete_empty_lines:
            doc1_paragraphs = [p for p in doc1_paragraphs if p.strip()]
            doc2_paragraphs = [p for p in doc2_paragraphs if p.strip()]
    else:
        # Unsupported file types or mismatch
        return 0

    if content_only:
        try:
            from rapidfuzz import fuzz
        except ImportError:
            fuzz = None
        if fuzz is None:
            return 0
        text1 = re.sub(r'\s+', ' ', '\n'.join(doc1_paragraphs)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        similarity = fuzz.ratio(text1, text2) / 100.0
        return similarity

    if ignore_blanks:
        text1 = re.sub(r'\s+', ' ', '\n'.join(doc1_paragraphs)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()

        if fuzzy_match:
            try:
                from rapidfuzz import fuzz
            except ImportError:
                return 0
            similarity = fuzz.ratio(text1, text2) / 100.0
            return similarity
        else:
            if text1 != text2:
                return 0
    else:
        if len(doc1_paragraphs) != len(doc2_paragraphs):
            return 0

        if fuzzy_match:
            try:
                from rapidfuzz import fuzz
            except ImportError:
                return 0
            total_similarity = 0
            if not doc1_paragraphs:
                return 1.0
            for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
                if ignore_case:
                    p1, p2 = p1.lower(), p2.lower()
                total_similarity += fuzz.ratio(p1, p2) / 100.0

            if len(doc1_paragraphs) == 0:
                return 1.0 if len(doc2_paragraphs) == 0 else 0.0

            avg_similarity = total_similarity / len(doc1_paragraphs)
            return avg_similarity
        else:
            for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
                if ignore_case:
                    p1, p2 = p1.lower(), p2.lower()
                if p1 != p2:
                    return 0

    return 1


# =============================================================================
# Metric 1: is_first_line_centered — 覆盖条件(1)标题居中
# 检查文档首段是否居中对齐
# =============================================================================

def is_first_line_centered(docx_file):
    if not docx_file:
        return 0

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    first_paragraph = doc.paragraphs[0]
    return 1 if first_paragraph.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER else 0


# =============================================================================
# Metric 2: check_title_font_size_20pt — 覆盖条件(1)标题20pt字号
# 检查文档首段首个 run 的字号是否为 20pt
# =============================================================================

def check_title_font_size_20pt(result_path, **options):
    if not result_path:
        return 0

    try:
        doc = Document(result_path)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    if not doc.paragraphs:
        return 0

    first_paragraph = doc.paragraphs[0]
    if not first_paragraph.runs:
        return 0

    first_run = first_paragraph.runs[0]
    if first_run.font.size is None:
        return 0

    return 1 if first_run.font.size.pt == 20 else 0


# =============================================================================
# Metric 3: check_attendees_bold — 覆盖条件(3)参会者列表段落加粗
# 以 gold 文档为参考，找出 gold 中加粗的段落，验证 result 中对应段落也已加粗
# =============================================================================

def check_attendees_bold(result_path, expected, **options):
    if not result_path or not expected:
        return 0

    try:
        result_doc = Document(result_path)
        gold_doc = Document(expected)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    found_bold_in_gold = False

    for i, gold_para in enumerate(gold_doc.paragraphs):
        if not gold_para.text.strip():
            continue
        gold_has_bold = any(run.bold for run in gold_para.runs if run.text.strip())
        if gold_has_bold:
            found_bold_in_gold = True
            if i >= len(result_doc.paragraphs):
                logger.error(f"Result doc has fewer paragraphs ({len(result_doc.paragraphs)}) "
                             f"than gold ({len(gold_doc.paragraphs)}), "
                             f"expected bold at index {i}")
                return 0
            result_para = result_doc.paragraphs[i]
            result_has_bold = any(run.bold for run in result_para.runs if run.text.strip())
            if not result_has_bold:
                logger.error(f"Paragraph {i} expected to be bold but is not")
                return 0

    if not found_bold_in_gold:
        logger.error("No bold paragraphs found in gold document")
        return 0

    return 1


# =============================================================================
# Metric 4: check_action_items_italic — 覆盖条件(4)行动项段落斜体
# 以 gold 文档为参考，找出 gold 中斜体的段落，验证 result 中对应段落也已斜体
# =============================================================================

def check_action_items_italic(result_path, expected, **options):
    if not result_path or not expected:
        return 0

    try:
        result_doc = Document(result_path)
        gold_doc = Document(expected)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    found_italic_in_gold = False

    for i, gold_para in enumerate(gold_doc.paragraphs):
        if not gold_para.text.strip():
            continue
        gold_has_italic = any(run.italic for run in gold_para.runs if run.text.strip())
        if gold_has_italic:
            found_italic_in_gold = True
            if i >= len(result_doc.paragraphs):
                logger.error(f"Result doc has fewer paragraphs ({len(result_doc.paragraphs)}) "
                             f"than gold ({len(gold_doc.paragraphs)}), "
                             f"expected italic at index {i}")
                return 0
            result_para = result_doc.paragraphs[i]
            result_has_italic = any(run.italic for run in result_para.runs if run.text.strip())
            if not result_has_italic:
                logger.error(f"Paragraph {i} expected to be italic but is not")
                return 0

    if not found_italic_in_gold:
        logger.error("No italic paragraphs found in gold document")
        return 0

    return 1


# =============================================================================
# Metric 5: check_section_headings_color — 覆盖条件(5)章节标题深蓝色 #003366
# 以 gold 文档为参考，找出 gold 中颜色为 target_color 的段落，
# 验证 result 中对应段落所有 run 的颜色也为 target_color
# =============================================================================

def check_section_headings_color(result_path, expected, **options):
    if not result_path or not expected:
        return 0

    target_color = options.get('target_color', '003366')

    try:
        result_doc = Document(result_path)
        gold_doc = Document(expected)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0

    found_target_color_in_gold = False

    for i, gold_para in enumerate(gold_doc.paragraphs):
        if not gold_para.text.strip():
            continue

        # Check if any run in this gold paragraph has the target color
        gold_has_target_color = False
        for run in gold_para.runs:
            if run.font.color and run.font.color.rgb:
                gold_color = str(run.font.color.rgb)
                if gold_color == target_color:
                    gold_has_target_color = True
                    break

        if gold_has_target_color:
            found_target_color_in_gold = True
            if i >= len(result_doc.paragraphs):
                logger.error(f"Result doc has fewer paragraphs ({len(result_doc.paragraphs)}) "
                             f"than gold ({len(gold_doc.paragraphs)}), "
                             f"expected color at index {i}")
                return 0

            result_para = result_doc.paragraphs[i]
            for run in result_para.runs:
                if run.text.strip():
                    if run.font.color is None or run.font.color.rgb is None:
                        logger.error(f"Paragraph {i} run has no color, expected #{target_color}")
                        return 0
                    result_color = str(run.font.color.rgb)
                    if result_color != target_color:
                        logger.error(f"Paragraph {i} run color is #{result_color}, "
                                     f"expected #{target_color}")
                        return 0

    if not found_target_color_in_gold:
        logger.error(f"No paragraphs with color #{target_color} found in gold document")
        return 0

    return 1
