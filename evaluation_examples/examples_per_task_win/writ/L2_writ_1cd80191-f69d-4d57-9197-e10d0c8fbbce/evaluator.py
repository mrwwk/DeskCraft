"""
Evaluator for L2_writ task: Policy_L2_04.docx formatting check.

Checks:
  1. All body text runs have font.size == 12pt (default, configurable via options.font_size_pt)
  2. Body text under section "5. Communication Standards" has italic=True on every run
  3. Body text NOT under section 5 has italic=None or italic=False (no spurious italic)
"""

import logging
import re

from docx import Document

logger = logging.getLogger("desktopenv.metric.docs")


def check_docx_formatting(result_path, expected_path=None, **options):
    """
    Verify formatting of Policy_L2_04.docx against task requirements.

    Args:
        result_path: Path to the agent-modified .docx file (from result getter).
        expected_path: Path to the gold .docx file (from expected getter), used
                       only for structural reference (identifying header paragraphs).
        **options:
            font_size_pt (float): Expected font size in points. Default 12.

    Returns:
        float: Score in [0.0, 1.0]. Proportional to how many of the three
               sub-checks (font size / section-5 italic / non-section-5 no-italic)
               pass.
    """
    if not result_path:
        logger.error("check_docx_formatting: result_path is empty")
        return 0.0

    try:
        doc = Document(result_path)
    except Exception as e:
        logger.error(f"check_docx_formatting: failed to open {result_path}: {e}")
        return 0.0

    font_size_pt = float(options.get("font_size_pt", 12))

    # ------------------------------------------------------------------
    # 1. Identify structural (non-body) paragraphs
    # ------------------------------------------------------------------
    section_header_re = re.compile(r'^\d+\.\s+')

    # Build set of header paragraph indices from the actual document
    header_indices = set()
    first_section_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if section_header_re.match(text):
            header_indices.add(i)
            if first_section_idx is None:
                first_section_idx = i

    # Paragraphs before the first numbered section header are
    # title / subtitle / metadata — treat as structural (not body).
    if first_section_idx is not None:
        for i in range(first_section_idx):
            if doc.paragraphs[i].text.strip():
                header_indices.add(i)

    # Also use the gold document (if available) to pick up any headers
    # that the agent may have mangled so they no longer match the regex.
    if expected_path:
        try:
            gold_doc = Document(expected_path)
            for gp in gold_doc.paragraphs:
                gtext = gp.text.strip()
                if not gtext:
                    continue
                if section_header_re.match(gtext):
                    # find matching paragraph in actual doc by text prefix
                    for i, para in enumerate(doc.paragraphs):
                        if para.text.strip().startswith(gtext[:20]):
                            header_indices.add(i)
                            break
        except Exception as e:
            logger.warning(f"check_docx_formatting: could not read gold doc: {e}")

    # ------------------------------------------------------------------
    # 2. Locate section 5 boundaries
    # ------------------------------------------------------------------
    section_5_start_idx = None
    section_6_start_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if section_header_re.match(text):
            if section_5_start_idx is None and (text.startswith("5.") or text.startswith("5 ")):
                section_5_start_idx = i
            elif section_5_start_idx is not None and section_6_start_idx is None:
                section_6_start_idx = i
                break

    # ------------------------------------------------------------------
    # 3. Check formatting on body paragraphs
    # ------------------------------------------------------------------
    font_size_ok = True
    section_5_italic_ok = True
    non_section_5_italic_ok = True
    body_paragraphs_checked = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if i in header_indices:
            continue

        body_paragraphs_checked += 1

        # Is this paragraph under section 5?
        in_section_5 = (
            section_5_start_idx is not None
            and i > section_5_start_idx
            and (section_6_start_idx is None or i < section_6_start_idx)
        )

        for run in para.runs:
            run_text = run.text
            if not run_text or not run_text.strip():
                continue

            # --- font size ---
            if run.font.size is None:
                font_size_ok = False
            elif abs(run.font.size.pt - font_size_pt) > 0.01:
                font_size_ok = False

            # --- italic ---
            if in_section_5:
                # Section 5 body: every run must be italic
                if not run.font.italic:
                    section_5_italic_ok = False
            else:
                # Non-section-5 body: no run should be italic
                if run.font.italic:
                    non_section_5_italic_ok = False

    if body_paragraphs_checked == 0:
        logger.warning("check_docx_formatting: no body paragraphs found to check")
        return 0.0

    # ------------------------------------------------------------------
    # 4. Compute proportional score
    # ------------------------------------------------------------------
    checks_passed = sum([font_size_ok, section_5_italic_ok, non_section_5_italic_ok])
    score = checks_passed / 3.0

    logger.info(
        "check_docx_formatting: font_size_ok=%s, section_5_italic_ok=%s, "
        "non_section_5_italic_ok=%s, body_paras=%d, score=%.2f",
        font_size_ok, section_5_italic_ok, non_section_5_italic_ok,
        body_paragraphs_checked, score,
    )

    return score
