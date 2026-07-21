# Dedicated evaluator for Policy_L3_04.docx task (cd8b1536)
# Replaces generic compare_docx_files with format-aware multi-metric checks.
# Conditions: (1) title center/22pt/bold, (2) section headings uppercase,
# (3) body font Times New Roman 11pt, (4) section 4 body bold,
# (5) first para CONFIDENTIAL, (6) last para Document Control.

import re
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from rapidfuzz import fuzz

logger = logging.getLogger("desktopenv.metric.docs")

_HEADING_RE = re.compile(r'^\d+\.\s+([A-Z][A-Z\s&/,\-]+)$')
_SPECIAL_MARKERS = [
    "CONFIDENTIAL",
    "Document Control",
    "COMPANY REMOTE WORK POLICY",
    "EFFECTIVE DATE",
]


def _normalize_dash(text: str) -> str:
    text = text.replace('\u2013', '-')
    text = text.replace('\u2014', '-')
    text = text.replace('\u2015', '-')
    return text


def _normalize_ws(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def _normalize_text(text: str) -> str:
    return _normalize_ws(_normalize_dash(text))


def _is_heading(text: str) -> bool:
    return bool(_HEADING_RE.match(text.strip()))


def _is_special(text: str) -> bool:
    lowered = text.lower()
    return any(marker.lower() in lowered for marker in _SPECIAL_MARKERS)


def _is_body_paragraph(text: str) -> bool:
    stripped = text.strip()
    return bool(stripped) and not _is_heading(stripped) and not _is_special(stripped)


def check_title_format(result_path, **options):
    if not result_path:
        return 0.0
    try:
        doc = Document(result_path)
    except Exception as exc:
        logger.error("check_title_format: cannot open %s: %s", result_path, exc)
        return 0.0

    title_keywords = options.get("title_keywords", "REMOTE WORK POLICY")
    expected_size_pt = float(options.get("title_font_size_pt", 22.0))

    for para in doc.paragraphs:
        if title_keywords.lower() not in para.text.lower():
            continue
        if para.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            return 0.0
        text_runs = [run for run in para.runs if run.text.strip()]
        if not text_runs:
            return 0.0
        if not any(run.bold for run in text_runs):
            return 0.0
        if any(run.font.size is not None and abs(run.font.size.pt - expected_size_pt) > 0.5 for run in text_runs):
            return 0.0
        return 1.0

    return 0.0


def check_body_font(result_path, **options):
    if not result_path:
        return 0.0
    try:
        doc = Document(result_path)
    except Exception as exc:
        logger.error("check_body_font: cannot open %s: %s", result_path, exc)
        return 0.0

    expected_font = options.get("expected_font", "Times New Roman")
    expected_size_pt = float(options.get("expected_size_pt", 11.0))
    threshold = float(options.get("body_font_threshold", 0.80))
    total_runs = 0
    failed_runs = 0

    for para in doc.paragraphs:
        if not _is_body_paragraph(para.text):
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            ok = True
            if run.font.name is not None and run.font.name != expected_font:
                ok = False
            if run.font.size is not None and abs(run.font.size.pt - expected_size_pt) > 0.5:
                ok = False
            if not ok:
                failed_runs += 1

    if total_runs == 0:
        return 1.0
    ratio = 1.0 - failed_runs / total_runs
    return 1.0 if ratio >= threshold else ratio


def check_section4_bold(result_path, **options):
    if not result_path:
        return 0.0
    try:
        doc = Document(result_path)
    except Exception as exc:
        logger.error("check_section4_bold: cannot open %s: %s", result_path, exc)
        return 0.0

    section4_marker = options.get("section4_marker", "4. EQUIPMENT AND SECURITY")
    sec4_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if section4_marker.lower() in para.text.lower():
            sec4_idx = idx
            break
    if sec4_idx is None:
        return 1.0

    next_sec_idx = None
    for idx in range(sec4_idx + 1, len(doc.paragraphs)):
        if _is_heading(doc.paragraphs[idx].text):
            next_sec_idx = idx
            break

    body_paras = [
        para for para in doc.paragraphs[sec4_idx + 1: next_sec_idx or len(doc.paragraphs)]
        if para.text.strip() and not _is_special(para.text)
    ]
    if not body_paras:
        return 1.0

    for para in body_paras:
        if not any(run.text.strip() and run.bold for run in para.runs):
            return 0.0
    return 1.0


def compare_docx_text_normalized(result_path, expected_path, **options):
    if not result_path or not expected_path:
        return 0.0
    try:
        result_doc = Document(result_path)
        expected_doc = Document(expected_path)
    except Exception as exc:
        logger.error("compare_docx_text_normalized: cannot open docs: %s", exc)
        return 0.0

    result_paras = [_normalize_text(p.text) for p in result_doc.paragraphs]
    result_paras = [text for text in result_paras if text]
    expected_paras = [_normalize_text(p.text) for p in expected_doc.paragraphs]
    expected_paras = [text for text in expected_paras if text]
    if not result_paras:
        return 0.0

    confidential_target = _normalize_text("CONFIDENTIAL - Internal Use Only")
    if result_paras[0] != confidential_target:
        return 0.0

    doc_control_prefix = _normalize_text("Document Control: Version 3.0")
    if doc_control_prefix not in result_paras[-1]:
        return 0.0

    for text in result_paras:
        match = _HEADING_RE.match(text)
        if match:
            heading_body = match.group(1)
            if heading_body != heading_body.upper():
                return 0.0

    similarity = fuzz.ratio('\n'.join(result_paras), '\n'.join(expected_paras)) / 100.0
    threshold = float(options.get("text_similarity_threshold", 0.85))
    return similarity if similarity >= threshold else 0.0
