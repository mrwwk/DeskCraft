import logging
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger("desktopenv.metric.docs")

# Known section heading texts for this proposal document
_SECTION_HEADINGS = {
    "Background", "Objectives", "Scope of Work", "Budget",
    "Expected Benefits", "Conclusion", "Recommendations",
    "Introduction", "Methodology", "Findings", "Analysis",
}


def compare_docx_content_and_format(result_path, expected_path=None, **options):
    """
    Compare a .docx result file against a gold reference for both text content
    and formatting.  Checks all six requirements from the task instruction:

    1. Title + "Submitted to" + "Date" lines are center-aligned.
    2. Title is 22 pt bold.
    3. All section headings are dark blue (#003366) and bold.
    4. All body text is 11 pt.
    5. The body paragraph immediately after "Expected Benefits" is italic.
    6. The document ends with the approval line
       "Approved by: _________________ Date: _________________".

    Parameters
    ----------
    result_path : str or None
        Path to the agent-produced .docx file.
    expected_path : str or None
        Path to the gold .docx file (used for text-content verification).
    options : dict
        Forwarded but currently unused.

    Returns
    -------
    float
        1.0 if all checks pass, 0.0 otherwise.
    """
    if not result_path:
        logger.error("result_path is missing")
        return 0.0

    # ------------------------------------------------------------------
    # 0. Load documents
    # ------------------------------------------------------------------
    try:
        result_doc = Document(result_path)
    except Exception as exc:
        logger.error(f"Cannot open result document: {exc}")
        return 0.0

    gold_doc = None
    if expected_path:
        try:
            gold_doc = Document(expected_path)
        except Exception as exc:
            logger.warning(f"Cannot open gold document (text check skipped): {exc}")

    result_paras = [p for p in result_doc.paragraphs if p.text.strip()]

    # ------------------------------------------------------------------
    # 1. Text-content check against gold (verifies no deletion + approval line)
    # ------------------------------------------------------------------
    if gold_doc is not None:
        gold_paras = [p for p in gold_doc.paragraphs if p.text.strip()]
        # Compare normalised paragraph texts one-to-one.
        # Allow trailing empty paragraphs to differ.
        min_len = min(len(result_paras), len(gold_paras))
        for idx in range(min_len):
            r_text = _normalize(result_paras[idx].text)
            g_text = _normalize(gold_paras[idx].text)
            if r_text != g_text:
                logger.error(
                    "Text mismatch at paragraph %d: "
                    "result=%r  gold=%r",
                    idx,
                    result_paras[idx].text[:120],
                    gold_paras[idx].text[:120],
                )
                return 0.0
        if len(result_paras) != len(gold_paras):
            logger.error(
                "Paragraph count mismatch: result=%d gold=%d",
                len(result_paras), len(gold_paras),
            )
            return 0.0

    # ------------------------------------------------------------------
    # 2. Identify key paragraphs by content heuristics
    # ------------------------------------------------------------------
    # The document is expected to have this structure:
    #   [0] Title
    #   [1] Submitted to: ...
    #   [2] Date: ...
    #   ... section headings and body paragraphs ...
    #   last: "Approved by: ..."

    if len(result_paras) < 4:
        logger.error("Document too short (%d paragraphs)", len(result_paras))
        return 0.0

    title_para = result_paras[0]
    submitted_para = result_paras[1]
    date_para = result_paras[2]

    # Collect section heading paragraphs (short lines whose text matches
    # known heading strings).
    heading_paras = []
    expected_benefits_body = None
    heading_indices = set()

    for i, para in enumerate(result_paras):
        text = para.text.strip()
        if text in _SECTION_HEADINGS:
            heading_paras.append(para)
            heading_indices.add(i)
            # The body paragraph for "Expected Benefits" is the next
            # non-empty paragraph.
            if text == "Expected Benefits":
                for j in range(i + 1, len(result_paras)):
                    if result_paras[j].text.strip():
                        expected_benefits_body = result_paras[j]
                        break

    # Body paragraphs: all non-empty paragraphs that are not title,
    # submitted-to, date, headings, or the approval line.
    body_paras = []
    for i, para in enumerate(result_paras):
        if i in (0, 1, 2):
            continue
        if i in heading_indices:
            continue
        body_paras.append(para)

    # Approval line is the last paragraph
    approval_para = result_paras[-1]

    # ------------------------------------------------------------------
    # 3. Requirement 1 — Center-align title + two lines below
    # ------------------------------------------------------------------
    for label, para in [
        ("title", title_para),
        ("Submitted to", submitted_para),
        ("Date", date_para),
    ]:
        if para.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            logger.error(
                "%s paragraph is not center-aligned (alignment=%s)",
                label, para.paragraph_format.alignment,
            )
            return 0.0

    # ------------------------------------------------------------------
    # 4. Requirement 2 — Title 22 pt bold
    # ------------------------------------------------------------------
    title_bold = False
    title_22pt = False
    for run in title_para.runs:
        if not run.text.strip():
            continue
        if run.bold:
            title_bold = True
        if run.font.size is not None:
            sz = run.font.size.pt
            if abs(sz - 22.0) < 1.5:   # allow ~1.5 pt tolerance
                title_22pt = True
    if not title_bold:
        logger.error("Title is not bold")
        return 0.0
    if not title_22pt:
        logger.error("Title font size is not ~22 pt")
        return 0.0

    # ------------------------------------------------------------------
    # 5. Requirement 3 — Section headings dark blue (#003366) + bold
    # ------------------------------------------------------------------
    for hp in heading_paras:
        heading_ok = False
        for run in hp.runs:
            if not run.text.strip():
                continue
            is_bold = run.bold
            color_hex = ""
            try:
                if run.font.color and run.font.color.rgb:
                    color_hex = str(run.font.color.rgb).upper()
            except Exception:
                color_hex = ""
            if is_bold and color_hex == "003366":
                heading_ok = True
                break
        if not heading_ok:
            logger.error(
                "Section heading %r is not bold+dark-blue", hp.text[:80]
            )
            return 0.0

    # ------------------------------------------------------------------
    # 6. Requirement 4 — Body text 11 pt
    # ------------------------------------------------------------------
    for bp in body_paras:
        for run in bp.runs:
            if not run.text.strip():
                continue
            if run.font.size is None:
                # Accept style-inherited size (often None for Normal style)
                continue
            sz = run.font.size.pt
            if abs(sz - 11.0) > 1.5:
                logger.error(
                    "Body text has font size %.1f instead of ~11 pt: %r",
                    sz, run.text[:80],
                )
                return 0.0

    # ------------------------------------------------------------------
    # 7. Requirement 5 — Expected Benefits body italic
    # ------------------------------------------------------------------
    if expected_benefits_body is None:
        logger.error("Could not locate body paragraph after 'Expected Benefits'")
        return 0.0

    eb_italic = False
    for run in expected_benefits_body.runs:
        if run.text.strip() and run.italic:
            eb_italic = True
            break
    if not eb_italic:
        logger.error("'Expected Benefits' body paragraph is not italic")
        return 0.0

    # ------------------------------------------------------------------
    # 8. Requirement 6 — Approval line at end
    # ------------------------------------------------------------------
    approval_text = _normalize(approval_para.text)
    if "approved by:" not in approval_text.lower():
        logger.error("Approval line missing: %r", approval_para.text[:120])
        return 0.0

    return 1.0


def _normalize(text: str) -> str:
    """Collapse all whitespace sequences into a single space and strip."""
    return re.sub(r"\s+", " ", text).strip()
